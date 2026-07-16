"""Generates an editable Word (.docx) version of the Part 2 classification
report. The rubric's final submission must be a PDF (export_submission.py's
build_report_pdf already produces that), but a docx lets the "findings /
comments" sections actually get hand-edited -- the PDF's auto-generated
comments are deliberately factual placeholders, not real analysis. Edit
this docx, then save/export it as PDF for the final submission.

Reuses export_submission.py's data loading and chart-plotting logic rather
than duplicating it, so the two reports never disagree with each other.
"""
import io
import sqlite3

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from export_submission import (
    DB_NAME, load_dataframe, plot_histogram, relabel_with_titles,
    CHARTED_PROJECT_TYPES, truncate_label, TABLE_LABEL_MAX_CHARS,
)
from isic_taxonomy import code_to_title_map

DOCX_NAME = "SQ26_Classification_Report.docx"
TABLE_STYLE = "Light Grid Accent 1"


def _fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = TABLE_STYLE
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return table


def _add_stat_strip(doc, stats):
    """stats: [(label, value), ...] -- a two-row table approximating a
    stat-tile strip (big number over a small caption)."""
    table = doc.add_table(rows=2, cols=len(stats))
    table.style = TABLE_STYLE
    for i, (label, value) in enumerate(stats):
        val_para = table.cell(0, i).paragraphs[0]
        val_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = val_para.add_run(str(value))
        run.bold = True
        run.font.size = Pt(18)

        lbl_para = table.cell(1, i).paragraphs[0]
        lbl_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl_run = lbl_para.add_run(label)
        lbl_run.font.size = Pt(9)
    doc.add_paragraph()


def get_file_class_counts(conn, repository_id, project_type):
    """FILES.class distribution -- the per-file classification (rubric
    Part 2 Step 3: classify each primary data file, not just the project)
    that classifier.py already writes to the database but that never
    appeared in the PDF/xlsx report until now."""
    query = """
    SELECT f.class FROM FILES f
    JOIN PROJECTS p ON f.project_id = p.id
    WHERE p.repository_id = ? AND p.project_type = ? AND f.class IS NOT NULL
    """
    file_df = pd.read_sql_query(query, conn, params=(repository_id, project_type))
    return file_df['class'].value_counts()


def build_overview(doc, conn, df):
    doc.add_heading("SQ26 Part 2 — QDArchive Classification Report", level=0)
    doc.add_paragraph(
        "Draft report — edit the “Findings” paragraphs and this overview "
        "before exporting to PDF for submission."
    )

    total_projects = len(df)
    total_files = pd.read_sql_query("SELECT COUNT(*) AS n FROM FILES", conn)["n"].iloc[0]
    classified_projects = df["primary_class"].notna().sum()
    classified_files = pd.read_sql_query(
        "SELECT COUNT(*) AS n FROM FILES WHERE class IS NOT NULL", conn
    )["n"].iloc[0]

    doc.add_heading("Executive Overview", level=1)
    _add_stat_strip(doc, [
        ("Repositories", df["repository_id"].nunique()),
        ("Total projects", total_projects),
        ("Total files", total_files),
        ("Projects classified", f"{classified_projects}/{total_projects}"),
        ("Files classified", f"{classified_files}/{total_files}"),
    ])

    type_counts = df["project_type"].value_counts()
    doc.add_heading("Project-Type Composition (all repositories)", level=2)
    _add_table(doc, ["Project Type", "Count"],
               [[t, int(type_counts.get(t, 0))] for t in
                ["QDA_PROJECT", "QD_PROJECT", "OTHER_PROJECT", "NOT_A_PROJECT"]])

    doc.add_paragraph(
        "Scope: this delivery covers the two assigned repositories -- 11 (Finnish Social Science "
        "Data Archive) and 20 (Sikt). No peer/shared databases from other students were imported "
        "into this scope; every number below comes directly from what this pipeline itself "
        "discovered and downloaded."
    )
    doc.add_page_break()


def build_repository_section(doc, conn, df, repo_id, titles):
    repo_df = df[df["repository_id"] == repo_id]
    doc.add_heading(f"Repository {repo_id}", level=1)

    type_counts = repo_df["project_type"].value_counts()
    _add_table(doc, ["Project Type", "Count"],
               [[t, int(type_counts.get(t, 0))] for t in
                ["QDA_PROJECT", "QD_PROJECT", "OTHER_PROJECT", "NOT_A_PROJECT"]])
    doc.add_paragraph()

    any_distribution = False
    for project_type in CHARTED_PROJECT_TYPES:
        type_df = repo_df[repo_df["project_type"] == project_type]
        class_counts = type_df["primary_class"].value_counts()
        if class_counts.empty:
            continue
        any_distribution = True

        doc.add_heading(f"{project_type} — by project", level=2)
        labeled = relabel_with_titles(class_counts, titles)
        fig, ax = plt.subplots(figsize=(9, max(4, len(labeled) * 0.35)))
        plot_histogram(ax, labeled, f"{project_type} — Repository {repo_id} (by project)")
        plt.tight_layout()
        doc.add_picture(_fig_to_stream(fig), width=Inches(6.3))

        top_20 = labeled.head(20)
        doc.add_heading(f"Top {len(top_20)} classes (by project)", level=3)
        _add_table(doc, ["Rank", "ISIC Class", "Count"],
                   [[i, truncate_label(c, TABLE_LABEL_MAX_CHARS), v]
                    for i, (c, v) in enumerate(top_20.items(), 1)])

        # Per-file distribution -- Part 2 Step 3 asks for this in addition
        # to the project-level view.
        file_counts = get_file_class_counts(conn, repo_id, project_type)
        if not file_counts.empty:
            file_labeled = relabel_with_titles(file_counts, titles)
            doc.add_heading(f"{project_type} — by primary file", level=2)
            fig2, ax2 = plt.subplots(figsize=(9, max(4, len(file_labeled) * 0.35)))
            plot_histogram(ax2, file_labeled, f"{project_type} — Repository {repo_id} (by file)")
            plt.tight_layout()
            doc.add_picture(_fig_to_stream(fig2), width=Inches(6.3))

        dominant = labeled.index[0]
        dominant_pct = class_counts.iloc[0] / len(type_df) if len(type_df) else 0
        p = doc.add_paragraph()
        p.add_run("Findings (draft — replace with your own interpretation): ").bold = True
        file_note = f", {len(file_counts)} distinct classes at the file level" if not file_counts.empty else ""
        p.add_run(
            f"{len(type_df)} {project_type} project(s) analyzed. Dominant class: {dominant} "
            f"({class_counts.iloc[0]}/{len(type_df)}, {dominant_pct:.0%}). "
            f"{len(class_counts)} distinct primary classes at the project level{file_note}."
        )
        doc.add_paragraph()

    if not any_distribution:
        doc.add_paragraph(
            "No QDA_PROJECT or QD_PROJECT records with an assigned primary_class were found "
            "for this repository, so no distribution is shown."
        )

    doc.add_page_break()


def build_limitations_section(doc):
    doc.add_heading("Technical Data Challenges and Limitations", level=1)
    doc.add_paragraph(
        "Data-availability and data-quality findings encountered while acquiring and classifying "
        "this dataset -- documented as data handling considerations, not code defects."
    )

    items = [
        ("No downloadable QDA analysis files (0 QDA_PROJECT)",
         "Every file on disk was checked against the full REFI-QDA/MaxQDA/NVivo/ATLAS.ti/QDA "
         "Miner/Quirkos/f4analyse extension list -- zero matches. The one real .qdpx file found "
         "on Sikt/Dataverse.no (a Nacey metaphor-analysis dataset) returns HTTP 403: its "
         "depositor restricted that specific file while leaving sibling files in the same study "
         "open. FSD's QDA-bearing datasets are Condition B -- released only after a human "
         "submits a real \"purpose of use\" application (project title, research description) "
         "per dataset. A login is necessary but not sufficient, and that application step was "
         "deliberately not automated."),
        ("Non-qualitative content mixed into the Sikt results",
         "The most common file extensions downloaded are .hsat-out, .hsat, .class, and .java -- "
         "software/tooling artifacts, not qualitative research data. At least one Sikt hit "
         "(\"navajo\") is a Java/Maven software repository, not a research dataset."),
        ("Some Sikt \"projects\" are individual loose files, not grouped studies",
         "Sikt/Dataverse.no's file-level search returns one hit per file rather than one hit per "
         "dataset. Files sharing no common filename prefix each became their own single-file "
         "\"project\" (e.g. fields-05600.dat). Their titles carry no identifiable topical content, "
         "so they correctly receive no ISIC classification rather than a forced, meaningless one."),
        ("Archive contents were initially invisible to classification",
         "Zip/tar archives were first recorded as one opaque file each. The primary data files "
         "bundled inside weren't individually indexed until archives were extracted in place and "
         "re-scanned."),
        ("Generic academic boilerplate initially skewed classification",
         "Words like \"research\", \"study\", and \"data\" are near-universal in project titles on "
         "this archive but rare across the 87 ISIC division reference texts, giving them "
         "inflated discriminating weight and pulling results toward division N72 (Scientific "
         "research and development) regardless of actual subject matter. Domain-specific "
         "stopwords were added to correct this."),
    ]
    for title, body in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title + " — ").bold = True
        p.add_run(body)
    doc.add_page_break()


def build_conclusion(doc, conn, df):
    doc.add_heading("Conclusion", level=1)
    total_projects = len(df)
    classified_projects = df["primary_class"].notna().sum()
    classified_files = pd.read_sql_query(
        "SELECT COUNT(*) AS n FROM FILES WHERE class IS NOT NULL", conn
    )["n"].iloc[0]
    doc.add_paragraph(
        f"Across the two assigned repositories, {total_projects} project records were reviewed; "
        f"{classified_projects} received a primary ISIC class and {classified_files} individual "
        f"primary/QDA files were classified at the file level. No QDA_PROJECT was identified in "
        f"either repository under the current search queries and access constraints -- the "
        f"Technical Data Challenges section above documents why, with direct evidence (HTTP "
        f"status codes, extension scans) rather than assumption."
    )


def main():
    conn = sqlite3.connect(DB_NAME, timeout=30.0)
    df = load_dataframe(conn)
    titles = code_to_title_map()

    doc = Document()
    build_overview(doc, conn, df)
    for repo_id in sorted(df["repository_id"].dropna().unique()):
        build_repository_section(doc, conn, df, int(repo_id), titles)
    build_limitations_section(doc)
    build_conclusion(doc, conn, df)

    doc.save(DOCX_NAME)
    conn.close()
    print(f"Saved editable draft report: {DOCX_NAME}")


if __name__ == "__main__":
    main()
